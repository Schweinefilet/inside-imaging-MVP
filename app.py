# app.py
import os
import io
import re
import json
import logging
import time
import click
from pathlib import Path
import boto3
from botocore.config import Config

# load .env early
from dotenv import load_dotenv
load_dotenv(dotenv_path=".env", override=True)

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    session,
    make_response,
    jsonify,
    abort,
    Response,
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_cors import CORS
from authlib.integrations.flask_client import OAuth

# local db + auth
from src import db
from src.auth import require_role
from content.blog_posts import BLOG_POSTS
from src.dicom_handler import parse_dicom_metadata, extract_frame_as_png, store_dicom, sanitize_dicom_bytes
from src.storage import get_storage

_AWS_REGION = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "us-east-1"

# Configure logging before any logging calls
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

# Optional: ship logs to AWS CloudWatch when AWS_CLOUDWATCH_LOG_GROUP is set.
# This gives you an immutable, off-host audit trail. Falls back silently when
# unset so local dev is unaffected.
_cw_group = os.environ.get("AWS_CLOUDWATCH_LOG_GROUP")
if _cw_group:
    try:
        import boto3 as _boto3_cw
        import watchtower
        _cw_client = _boto3_cw.client(
            "logs",
            region_name=os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "us-east-1",
        )
        _cw_handler = watchtower.CloudWatchLogHandler(
            log_group_name=_cw_group,
            stream_name=os.environ.get("AWS_CLOUDWATCH_LOG_STREAM", "app"),
            boto3_client=_cw_client,
            send_interval=5,
            create_log_group=True,
        )
        _cw_handler.setLevel(logging.INFO)
        _cw_handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(name)s %(message)s"))
        logging.getLogger().addHandler(_cw_handler)
        logging.info("CloudWatch log shipping enabled (group=%s)", _cw_group)
    except Exception:
        logging.exception("CloudWatch log shipping requested but failed to initialize; continuing without it")

logging.info("INSIDEIMAGING_ALLOW_LLM=%r", os.getenv("INSIDEIMAGING_ALLOW_LLM"))
_OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5")
logging.info("OPENAI_MODEL=%r", _OPENAI_MODEL)

# Lazy-loading singleton for the AWS Textract client
_textract_client = None

def _textract():
    global _textract_client
    if _textract_client is None:
        _textract_client = boto3.client(
            "textract",
            region_name=_AWS_REGION,
            config=Config(retries={"max_attempts": 3, "mode": "standard"})
        )
    return _textract_client

# Helpers for tenant-scoped DICOM routes
def _current_tenant():
    return db.get_user_tenant(session.get("username", ""))


def _audit(action, **kwargs):
    try:
        db.log_audit_event(
            tenant_id=_current_tenant(),
            username=session.get("username", ""),
            action=action,
            ip_address=request.remote_addr or "",
            user_agent=request.headers.get("User-Agent", ""),
            **kwargs,
        )
    except Exception:
        logging.exception("Audit log write failed (action=%s)", action)

# --- app ---
app = Flask(__name__)

# Phase 0.7: production detection
_ENV = os.environ.get("INSIDEIMAGING_ENV", os.environ.get("FLASK_ENV", "development")).lower()
_IS_PRODUCTION = _ENV in ("production", "prod")

# Phase 0.1: SECRET_KEY — fail loudly in production, warn in dev
_secret_key = os.environ.get("SECRET_KEY")
if not _secret_key:
    if _IS_PRODUCTION:
        raise SystemExit(
            "FATAL: SECRET_KEY env var is required in production. Refusing to start. "
            "Generate one with: python -c 'import secrets; print(secrets.token_urlsafe(64))'"
        )
    logging.warning("SECRET_KEY env var not set — using insecure default. Set SECRET_KEY in production.")
    _secret_key = "dev-only-do-not-use-in-production"
app.secret_key = _secret_key

# Phase 0.2: session cookies + general hardening
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=_IS_PRODUCTION,
    PERMANENT_SESSION_LIFETIME=int(os.getenv("SESSION_LIFETIME_SECONDS", "43200")),  # 12h default
    MAX_CONTENT_LENGTH=int(os.getenv("MAX_UPLOAD_MB", "50")) * 1024 * 1024,
)

# Phase 0.5: CSRF protection (Flask-WTF CSRFProtect, no WTForms dependency required)
try:
    from flask_wtf.csrf import CSRFProtect, CSRFError
    csrf = CSRFProtect(app)

    @app.errorhandler(CSRFError)
    def _csrf_error(e):
        logging.warning("CSRF rejected: %s ip=%s ep=%s", e.description, request.remote_addr, request.endpoint)
        return jsonify({"error": "CSRF token missing or invalid"}), 400
except Exception:
    logging.exception("Flask-WTF not available; CSRF protection NOT active. Install Flask-WTF.")
    csrf = None

# Phase 1.1: rate limiting (login, integration endpoints)
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    limiter = Limiter(
        get_remote_address,
        app=app,
        default_limits=[os.getenv("RATELIMIT_DEFAULT", "300 per hour")],
        storage_uri=os.getenv("RATELIMIT_STORAGE", "memory://"),
        strategy="fixed-window",
    )
except Exception:
    logging.exception("Flask-Limiter not available; rate limiting NOT active. Install Flask-Limiter.")
    limiter = None

from src.integration import extensions as integration_extensions
integration_extensions.limiter = limiter

# OAuth configuration
oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=os.getenv('GOOGLE_CLIENT_ID'),
    client_secret=os.getenv('GOOGLE_CLIENT_SECRET'),
    access_token_url='https://accounts.google.com/o/oauth2/token',
    access_token_params=None,
    authorize_url='https://accounts.google.com/o/oauth2/auth',
    authorize_params=None,
    api_base_url='https://www.googleapis.com/oauth2/v1/',
    userinfo_endpoint='https://openidconnect.googleapis.com/v1/userinfo',
    client_kwargs={'scope': 'openid email profile'},
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration'
)

# CORS — CORS_ORIGINS can be a comma-separated list of allowed origins
_cors_origins_raw = os.getenv("CORS_ORIGINS", "")
_cors_origins = [o.strip() for o in _cors_origins_raw.split(",") if o.strip()] or ["https://schweinefilet.github.io"]
CORS(app, resources={r"/*": {"origins": _cors_origins}})

# Import the integration blueprint only after shared extensions are initialized.
from src.integration.api import integration_bp

# Integration blueprint is API-only (hospital systems call it with API key, not session-bound).
# CSRF on these endpoints would force hospital systems to fetch tokens; exempt the whole blueprint.
if csrf is not None:
    csrf.exempt(integration_bp)
app.register_blueprint(integration_bp)


# Phase 0.3b: session inactivity timeout
_INACTIVITY_TIMEOUT_SECONDS = int(os.getenv("SESSION_INACTIVITY_SECONDS", "1800"))  # 30 min


@app.before_request
def _check_inactivity_timeout():
    if "username" not in session:
        return None
    last = session.get("_last_active")
    now = time.time()
    if last and (now - float(last)) > _INACTIVITY_TIMEOUT_SECONDS:
        session.clear()
        if request.accept_mimetypes.accept_json and not request.accept_mimetypes.accept_html:
            return jsonify({"error": "Session expired"}), 401
        flash("Your session expired due to inactivity. Please log in again.", "info")
        # Don't force-redirect — let each route decide. Public pages (/, /dashboard,
        # /blogs) keep working; protected routes redirect to login themselves.
    session["_last_active"] = now
    return None


# Phase 0.3: force HTTPS in production
@app.before_request
def _force_https():
    if not _IS_PRODUCTION:
        return None
    # Hosts that terminate TLS upstream set X-Forwarded-Proto=https
    proto = request.headers.get("X-Forwarded-Proto", request.scheme)
    if proto != "https":
        url = request.url.replace("http://", "https://", 1)
        return redirect(url, code=301)
    return None


# Phase 0.4: security headers on every response
_CSP_DEFAULT = (
    "default-src 'self'; "
    "img-src 'self' data: blob: https:; "
    # 'unsafe-inline' kept until inline <style>/<script> blocks are fully extracted (Tier 3 of the redesign).
    "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
    "font-src 'self' https://fonts.gstatic.com data:; "
    "script-src 'self' 'unsafe-inline'; "
    "connect-src 'self'; "
    "frame-ancestors 'none'; "
    "base-uri 'self'; "
    "form-action 'self' https://accounts.google.com"
)
_CSP = os.getenv("CONTENT_SECURITY_POLICY", _CSP_DEFAULT)


@app.after_request
def _security_headers(response):
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    response.headers.setdefault("Permissions-Policy", "geolocation=(), camera=(), microphone=()")
    response.headers.setdefault("Cross-Origin-Opener-Policy", "same-origin")
    if _IS_PRODUCTION:
        response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    response.headers.setdefault("Content-Security-Policy", _CSP)
    return response


# Prevent browsers caching HTML responses (avoids stale theme scripts)
@app.after_request
def _no_cache_html(response):
    if response.content_type and "text/html" in response.content_type:
        response.headers["Cache-Control"] = "no-store"
    return response


# available languages
LANGUAGES = ["English", "Kiswahili"]

# pricing + tokens
USD_PER_REPORT = float(os.getenv("PRICE_USD_PER_REPORT", "1.00"))
KES_PER_USD = float(os.getenv("PRICE_KES_PER_USD", "129"))
TOKENS_PER_REPORT = int(os.getenv("TOKENS_PER_REPORT", "1"))

# curated content for magazine + blog pages
MAGAZINE_ISSUES = [
    {
        "title": "July 2025 · THE FUTURE OF AI IN IMAGING",
        "url": "magazine/July-2025.pdf",
        "note": "Upload static/magazine/July-2025.pdf",
    },
]

MARQUEE_IMAGES = [
    # Real radiology examples from the team
    "/static/images/marquee/IMG-20251030-WA0002.jpg",
    "/static/images/marquee/IMG-20251030-WA0003.jpg",
    "/static/images/marquee/IMG-20251030-WA0004.jpg",
    "/static/images/marquee/IMG-20251030-WA0005.jpg",
    "/static/images/marquee/IMG-20251030-WA0006.jpg",
    "/static/images/marquee/IMG-20251030-WA0007.jpg",
    "/static/images/marquee/IMG-20251030-WA0008.jpg",
    "/static/images/marquee/IMG-20251030-WA0009.jpg",
    "/static/images/marquee/IMG-20251030-WA0010.jpg",
    "/static/images/marquee/IMG-20251030-WA0011.jpg",
    "/static/images/marquee/IMG-20251030-WA0012.jpg",
    "/static/images/marquee/IMG-20251030-WA0013.jpg",
    "/static/images/marquee/IMG-20251030-WA0014.jpg",
    "/static/images/marquee/IMG-20251030-WA0015.jpg",
    "/static/images/marquee/IMG-20251030-WA0016.jpg",
    "/static/images/marquee/IMG-20251030-WA0017.jpg",
    "/static/images/marquee/IMG-20251030-WA0018.jpg",
    "/static/images/marquee/IMG-20251030-WA0019.jpg",
    "/static/images/marquee/IMG-20251030-WA0020.jpg",
]

# Initialize database
try:
    db.init_db()
except Exception:
    logging.exception("Database initialization failed")


@app.cli.command("backup-db")
@click.argument("label", default="")
def backup_db_command(label: str) -> None:
    """Create a timestamped backup of the database. Usage: flask backup-db [label]"""
    path = db.backup_db(label=label)
    if path:
        click.echo(f"Backup written to {path}")
    else:
        click.echo("No database file found — nothing to back up.")

# --- translate wiring ---
try:
    from src.translate import Glossary, build_structured  # type: ignore
    from src.translate.flag import check_report_sensitivity  # type: ignore
    from src.parse import parse_metadata  # type: ignore
except Exception:
    logging.exception("translate import failed")
    Glossary = None  # type: ignore
    parse_metadata = None  # type: ignore

    def check_report_sensitivity(text: str):  # type: ignore
        return {"flagged": False}

    def build_structured(report_text: str, glossary=None, language: str = "English"):
        return {
            "reason": "",
            "technique": "",
            "findings": (report_text or "").strip(),
            "conclusion": "",
            "concern": "",
        }

# try to load a glossary if you have one; otherwise None is fine
LAY_GLOSS = None
try:
    if Glossary:
        gloss_path = str(Path(__file__).parent / "data" / "glossary.csv")
        if os.path.exists(gloss_path):
            LAY_GLOSS = Glossary.load(gloss_path)
except Exception:
    logging.exception("glossary load failed")
    LAY_GLOSS = None

# PDF engine
try:
    from weasyprint import HTML  # type: ignore
except Exception:
    HTML = None  # type: ignore


def _extract_text_from_pdf_bytes(data: bytes) -> str:
    """Robust PDF text extraction using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text  # type: ignore
    except Exception:
        logging.exception("pdfminer.six not available")
        return ""
    try:
        return extract_text(io.BytesIO(data)) or ""
    except Exception:
        logging.exception("pdfminer extract_text failed")
        return ""


def _extract_text_from_image_bytes(data: bytes) -> str:
    """Extract text from images (JPEG/PNG).

    Works best for phone photos. Max 5 MB for Bytes input.
    """
    if len(data) > 5 * 1024 * 1024:
        logging.warning("Image >5MB; Textract DetectDocumentText requires <=5MB for Bytes.")
        return ""

    try:
        resp = _textract().detect_document_text(Document={"Bytes": data})
    except Exception:
        logging.exception("Textract DetectDocumentText failed")
        return ""

    # Pull out LINE blocks in natural reading order
    lines = []
    for block in resp.get("Blocks", []):
        if block.get("BlockType") == "LINE" and block.get("Text"):
            lines.append(block["Text"].strip())

    # Fallback: if no LINEs, try WORDs
    if not lines:
        words = [b.get("Text", "").strip() for b in resp.get("Blocks", []) if b.get("BlockType") == "WORD"]
        lines = [" ".join(w for w in words if w)]

    text = "\n".join(l for l in lines if l)
    return text.strip()


def _extract_text_from_docx_bytes(data: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document  # type: ignore
    except Exception:
        logging.exception("python-docx not available")
        return ""
    
    try:
        doc = Document(io.BytesIO(data))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        return "\n".join(paragraphs)
    except Exception:
        logging.exception("docx extraction failed")
        return ""


def _extract_text_from_heif_bytes(data: bytes) -> str:
    """
    Extract text from HEIF/HEIC images (iOS photos) by converting to JPEG
    and then using AWS Textract.
    """
    try:
        import pillow_heif  # type: ignore
        from PIL import Image  # type: ignore
        pillow_heif.register_heif_opener()
    except Exception:
        logging.exception("pillow-heif not available")
        return ""
    
    try:
        # Convert HEIF to JPEG in memory
        heif_image = pillow_heif.read_heif(io.BytesIO(data))
        if heif_image.data is None:
            raise ValueError("HEIF image has no pixel data")
        image = Image.frombytes(
            heif_image.mode,
            heif_image.size,
            heif_image.data,
            "raw",
        )
        
        # Convert to JPEG bytes
        jpeg_buffer = io.BytesIO()
        image.save(jpeg_buffer, format="JPEG", quality=95)
        jpeg_bytes = jpeg_buffer.getvalue()
        
        # Use existing image extraction with Textract
        return _extract_text_from_image_bytes(jpeg_bytes)
    except Exception:
        logging.exception("HEIF extraction failed")
        return ""




def _pdf_response_from_html(html_str: str, *, filename="inside-imaging-report.pdf", inline: bool = False):
    if not HTML:
        raise RuntimeError("WeasyPrint is not installed or failed to import")
    # host_url lets WeasyPrint resolve /static and relative asset URLs
    pdf_bytes = HTML(string=html_str, base_url=request.host_url).write_pdf()
    resp = make_response(pdf_bytes)
    resp.headers["Content-Type"] = "application/pdf"
    disp = "inline" if inline else "attachment"
    resp.headers["Content-Disposition"] = f'{disp}; filename="{filename}"'
    return resp




_TRIAGE_SECTION_RX = re.compile(
    r"(?im)^\s*(findings|impression|conclusion|technique|history|clinical\s+history|"
    r"indication|comparison|procedure|exam(?:ination)?|study|details)\s*[:\-]"
)
_TRIAGE_MODALITY_TOKENS = [
    "ct", "mri", "x-ray", "xray", "ultrasound", "pet", "spect", "angiogram",
    "fluoroscopy", "mammo", "mammogram", "cect", "mra", "cta", "doppler",
]
_TRIAGE_IMAGING_TERMS = [
    "lesion", "mass", "nodule", "enhancement", "attenuation", "hyperdense",
    "hypodense", "hyperintense", "hypointense", "density", "signal", "axial",
    "sagittal", "coronal", "sequence", "cm", "mm", "vertebra", "lobar",
    "hepatic", "renal", "ventricle", "parenchyma", "impression", "findings",
    "technique", "study", "comparison", "contrast",
]
_TRIAGE_NEGATIVE_TOKENS = [
    "syllabus", "semester", "homework", "assignment", "professor", "student",
    "lecture", "quiz", "final exam", "midterm", "credit hours", "office hours",
    "course objectives", "course description", "grading policy", "title ix",
    "canvas site", "attendance policy",
]


def _triage_radiology_report(text: str) -> tuple[bool, dict]:
    """Quick heuristic to reject non-radiology uploads before hitting the LLM."""

    sample = (text or "").strip()
    if not sample:
        return False, {"reason": "empty"}

    snippet = sample[:20000]
    lower = snippet.lower()

    # Basic counts
    words = re.findall(r"\b\w+\b", snippet)
    word_count = len(words)
    section_hits = {match.group(1).lower() for match in _TRIAGE_SECTION_RX.finditer(snippet)}
    modality_hits = [token for token in _TRIAGE_MODALITY_TOKENS if token in lower]
    imaging_hits = [token for token in _TRIAGE_IMAGING_TERMS if token in lower]
    measurement_count = len(re.findall(r"\b\d+(?:\.\d+)?\s*(?:mm|cm)\b", lower))
    negative_hits = [token for token in _TRIAGE_NEGATIVE_TOKENS if token in lower]

    # Legacy keyword heuristics to preserve prior thresholds
    radiology_keywords = [
        "radiology", "radiologist", "imaging", "scan", "ct", "mri", "x-ray", "xray",
        "ultrasound", "pet", "findings", "impression", "technique", "contrast",
        "examination", "study", "patient", "indication", "conclusion", "comparison",
    ]
    anatomy_terms = [
        "brain", "lung", "liver", "kidney", "heart", "spine", "abdomen", "pelvis",
        "chest", "thorax", "head", "skull", "bone", "soft tissue", "vessel", "artery",
        "vein", "organ", "lesion", "mass", "nodule",
    ]
    radiology_keyword_count = sum(1 for keyword in radiology_keywords if keyword in lower)
    anatomy_count = sum(1 for term in anatomy_terms if term in lower)

    score = 0
    if word_count >= 90:
        score += 1
    if len(section_hits) >= 2:
        score += 2
    elif len(section_hits) == 1:
        score += 1
    if radiology_keyword_count >= 3:
        score += 1
    if anatomy_count >= 2 or len(imaging_hits) >= 4:
        score += 1
    if modality_hits:
        score += 2
    if measurement_count >= 3:
        score += 2
    elif measurement_count >= 1:
        score += 1
    if "impression" in section_hits:
        score += 1
    if "findings" in section_hits:
        score += 1

    diagnostics = {
        "word_count": word_count,
        "sections": sorted(section_hits),
        "modalities": modality_hits,
        "imaging_hits": imaging_hits[:10],
        "radiology_keyword_count": radiology_keyword_count,
        "anatomy_count": anatomy_count,
        "measurement_count": measurement_count,
        "negative_hits": negative_hits,
        "score": score,
    }

    # Hard rejection conditions
    if word_count < 80 and not (len(section_hits) >= 3 and modality_hits):
        diagnostics["reason"] = "too_short"
        return False, diagnostics
    if negative_hits and score < 6:
        diagnostics["reason"] = "non_medical_tokens"
        return False, diagnostics
    if not modality_hits and len(section_hits) < 2 and len(imaging_hits) < 5:
        diagnostics["reason"] = "insufficient_radiology_markers"
        return False, diagnostics
    if score < 5:
        diagnostics["reason"] = "low_confidence"
        return False, diagnostics

    diagnostics["reason"] = "ok"
    return True, diagnostics



@app.route("/dashboard", methods=["GET"])
def dashboard():
    stats = db.get_stats()
    recent_reports = session.get("recent_reports", [])
    
    # Get user's persistent reports if logged in
    username = session.get("username")
    user_reports = []
    if username:
        try:
            user_reports = db.get_user_reports(username, limit=5)
        except Exception:
            logging.exception("Failed to fetch user reports")
    
    return render_template("index.html", stats=stats, languages=LANGUAGES, 
                          recent_reports=recent_reports, user_reports=user_reports)


# Dispatch table for file uploads: extension → (extractor, kind, flash-message-on-empty)
_IMG_ERR = ("Unable to extract text from the image. Please try a clearer image "
            "or paste the text directly.")
_UPLOAD_EXTRACTORS = {
    ".pdf":  (_extract_text_from_pdf_bytes,  "pdf",
              "Unable to extract text from the PDF. Please try a different file or paste the text directly."),
    ".heic": (_extract_text_from_heif_bytes, "heif",
              "Unable to extract text from the HEIF/HEIC image. Please try a different format or paste the text directly."),
    ".heif": (_extract_text_from_heif_bytes, "heif",
              "Unable to extract text from the HEIF/HEIC image. Please try a different format or paste the text directly."),
    ".docx": (_extract_text_from_docx_bytes, "docx",
              "Unable to extract text from the Word document. Please try a different file or paste the text directly."),
    ".png":  (_extract_text_from_image_bytes, "image", _IMG_ERR),
    ".jpg":  (_extract_text_from_image_bytes, "image", _IMG_ERR),
    ".jpeg": (_extract_text_from_image_bytes, "image", _IMG_ERR),
    ".webp": (_extract_text_from_image_bytes, "image", _IMG_ERR),
    ".tif":  (_extract_text_from_image_bytes, "image", _IMG_ERR),
    ".tiff": (_extract_text_from_image_bytes, "image", _IMG_ERR),
    ".bmp":  (_extract_text_from_image_bytes, "image", _IMG_ERR),
}


def _extract_uploaded_report(file, file_text):
    """Return (text, kind, error_message). A non-empty error triggers flash+redirect."""
    if file_text and file_text.strip():
        return file_text.strip(), "text", ""
    if not (file and file.filename):
        return "", "", ""

    fname = secure_filename(file.filename)
    ext = "." + fname.rsplit(".", 1)[-1].lower() if "." in fname else ""

    # For Textract-bound image types, check file size before reading the whole
    # file into memory — Textract rejects anything over 5 MB anyway.
    _IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}
    if ext in _IMAGE_EXTS:
        try:
            file.stream.seek(0, 2)
            size = file.stream.tell()
            file.stream.seek(0)
            if size > 5 * 1024 * 1024:
                return "", "image", (
                    "Image exceeds the 5 MB limit for text extraction. "
                    "Please compress it and try again."
                )
        except Exception:
            pass  # stream not seekable; fallback check inside extractor

    data = file.read()

    try:
        if ext in _UPLOAD_EXTRACTORS:
            extractor, kind, err = _UPLOAD_EXTRACTORS[ext]
            text = extractor(data) or ""
            return text, kind, ("" if text else err)
        # Unknown extension: best-effort UTF-8 decode
        return data.decode("utf-8", "ignore"), "text", ""
    except Exception:
        logging.exception("file handling failed; extracted empty")
        return "", ext.lstrip("."), ""


def _build_patient_block(structured, full_patient_name):
    """Patient display dict; restores full_patient_name (display-only, session-only)."""
    pstruct = structured.get("patient") if isinstance(structured, dict) else None
    src = pstruct if isinstance(pstruct, dict) and pstruct else structured
    return {
        "hospital": src.get("hospital", ""),
        "study":    src.get("study", "Unknown"),
        "name":     full_patient_name,
        "sex":      src.get("sex", ""),
        "age":      src.get("age", ""),
        "date":     src.get("date", ""),
        "history":  src.get("history", ""),
    }


def _persist_recent_report(patient, structured, report_stats, lang, context):
    """Persist analytics + prepend the new report to the session's recent list."""
    try:
        username = session.get("username", "")
        report_id = db.store_report_event(patient, structured, report_stats, lang, username, context)
    except Exception:
        logging.exception("Failed to persist report analytics.")
        return
    if not report_id:
        return
    try:
        brief = db.get_report_brief(report_id)
    except Exception:
        logging.exception("Failed to fetch report brief.")
        return
    if brief:
        history = [item for item in (session.get("recent_reports") or []) if item.get("id") != report_id]
        session["recent_reports"] = [brief] + history[:4]


@app.route("/upload", methods=["GET", "POST"])
def upload():
    if request.method == "GET":
        return redirect(url_for("dashboard"))

    lang = request.form.get("language", "English")
    context = request.form.get("context", "")

    extracted, kind, err = _extract_uploaded_report(
        request.files.get("file"),
        request.form.get("file_text", ""),
    )
    if err:
        flash(err, "error")
        return redirect(url_for("dashboard"))
    logging.info("len(extracted)=%s kind=%s", len(extracted), kind or "?")

    triage_ok, triage_diag = _triage_radiology_report(extracted)
    if not triage_ok:
        flash(
            "The uploaded file doesn't appear to be a radiology report. "
            "Please upload a full imaging report (with sections like Findings and Impression).",
            "error",
        )
        logging.warning("Upload triage rejected (reason=%s, words=%s)",
                        triage_diag.get("reason"), triage_diag.get("word_count"))
        return redirect(url_for("dashboard"))

    is_flagged = check_report_sensitivity(extracted)["flagged"]
    if is_flagged:
        logging.info(json.dumps({"event": "sensitivity_flag", "timestamp": time.time(), "flagged": True, "acknowledged": False}))

    # PHI handling: extract the full name from raw text BEFORE build_structured strips it,
    # then keep it ONLY in session (never sent to OpenAI). Display-only on result page.
    full_patient_name = ""
    if parse_metadata:
        try:
            full_patient_name = (parse_metadata(extracted) or {}).get("name", "")
            logging.info("Extracted patient name (length=%d) for session-only storage", len(full_patient_name))
        except Exception:
            logging.exception("Failed to parse metadata for name extraction")

    try:
        structured = build_structured(extracted, LAY_GLOSS, language=lang) or {}
    except Exception:
        logging.exception("build_structured failed")
        structured = {"reason": "", "technique": "", "findings": "", "conclusion": "", "concern": ""}

    patient = _build_patient_block(structured, full_patient_name)
    high_html = (structured.get("findings", "") or "") + (structured.get("conclusion", "") or "")
    report_stats = {
        "words": len(extracted.split()),
        "sentences": len(re.findall(r"[.!?]+", extracted)),
        "highlights_positive": high_html.count('class="ii-pos"'),
        "highlights_negative": high_html.count('class="ii-neg"'),
    }
    disease_tags = db.detect_disease_tags(
        (structured.get("findings") or "") + (structured.get("conclusion") or "") + (structured.get("concern") or "")
    )

    session["structured"] = structured
    session["patient"] = patient
    session["language"] = lang
    session["context"] = context

    _persist_recent_report(patient, structured, report_stats, lang, context)

    if context:
        structured["reason"] = f"<strong>Patient context:</strong> {context}<br><br>" + (structured.get("reason") or "")

    return render_template(
        "result.html",
        S=structured, structured=structured, patient=patient,
        extracted=extracted, study={"organ": patient.get("study") or "Unknown"},
        language=lang, report_stats=report_stats, disease_tags=disease_tags,
        flagged=is_flagged,
    )


@app.route("/flag-acknowledge", methods=["POST"])
def flag_acknowledge():
    logging.info(json.dumps({"event": "sensitivity_flag", "timestamp": time.time(), "flagged": True, "acknowledged": True}))
    return "", 204


@app.route("/reports/<int:report_id>")
def report_detail(report_id: int):
    if "username" not in session:
        return redirect(url_for("login"))
    record = db.get_report_detail(report_id)
    if not record:
        abort(404)
        return  # unreachable; narrows type past here
    owner = record["username"]
    if owner and owner != session["username"]:
        abort(403)

    structured = dict(record.get("structured") or {})
    patient = dict(record.get("patient") or {})
    language = record.get("language") or "English"

    findings_blob = (structured.get("findings") or "") + (structured.get("conclusion") or "")
    highlight_pos = findings_blob.count('class="ii-pos"')
    highlight_neg = findings_blob.count('class="ii-neg"')

    structured.setdefault("word_count", record.get("word_count", 0))
    structured.setdefault("sentence_count", 0)
    structured.setdefault("highlights_positive", highlight_pos)
    structured.setdefault("highlights_negative", highlight_neg)

    report_stats = {
        "words": structured.get("word_count", 0),
        "sentences": structured.get("sentence_count", 0),
        "highlights_positive": highlight_pos,
        "highlights_negative": highlight_neg,
    }

    session["structured"] = structured
    session["patient"] = patient
    session["language"] = language

    study = {"organ": patient.get("study") or "Unknown"}
    disease_tags = record.get("disease_tags") or []
    if isinstance(disease_tags, str) and disease_tags:
        disease_tags = [t.strip() for t in disease_tags.split(",") if t.strip()]

    return render_template(
        "result.html",
        S=structured,
        structured=structured,
        patient=patient,
        extracted="",
        study=study,
        language=language,
        report_stats=report_stats,
        disease_tags=disease_tags,
    )


@app.route("/download-pdf", methods=["GET", "POST"])
def download_pdf():
    try:
        if request.method == "POST":
            if request.is_json:
                body = request.get_json(silent=True) or {}
                structured = body["structured"] if "structured" in body else session.get("structured", {}) or {}
                patient = body["patient"] if "patient" in body else session.get("patient", {}) or {}
            else:
                structured_raw = request.form.get("structured")
                patient_raw = request.form.get("patient")
                structured = json.loads(structured_raw) if structured_raw else session.get("structured", {}) or {}
                patient = json.loads(patient_raw) if patient_raw else session.get("patient", {}) or {}
        else:
            structured = session.get("structured", {}) or {}
            patient = session.get("patient", {}) or {}
    except Exception:
        logging.exception("Failed to parse form JSON")
        return jsonify({"error": "bad form JSON"}), 400

    html_str = render_template("pdf_report.html", structured=structured, patient=patient)

    # hard fail if PDF fails. no HTML fallback.
    try:
        return _pdf_response_from_html(html_str, filename="inside-imaging-report.pdf", inline=False)
    except Exception:
        logging.exception("WeasyPrint PDF render failed")
        return jsonify({"error": "pdf_failed"}), 500


@app.get("/pdf-smoke")
def pdf_smoke():
    test_html = """
    <!doctype html><meta charset="utf-8">
    <style>@page{size:A4;margin:20mm} body{font-family:Arial}</style>
    <h1>WeasyPrint OK</h1><p>Static image test below.</p>
    <img src="/static/logo.png" alt="logo" height="24">
    """
    try:
        return _pdf_response_from_html(test_html, filename="smoke.pdf", inline=True)
    except Exception:
        logging.exception("Smoke failed")
        return jsonify({"error": "smoke_failed"}), 500


@app.get("/report/preview")
def report_preview():
    """Quick HTML preview of the PDF template with session data."""
    structured = session.get("structured", {}) or {}
    patient = session.get("patient", {}) or {}
    return render_template("pdf_report.html", structured=structured, patient=patient)


@app.route("/", methods=["GET"])
@app.route("/projects")
def projects():
    stats = db.get_stats()
    return render_template(
        "projects.html",
        posts=BLOG_POSTS,
        marquee_images=MARQUEE_IMAGES,
        submit_url="mailto:editor@insideimaging.example?subject=Radiologist%20Blog%20Pitch",
        stats=stats,
        languages=LANGUAGES,
    )


@app.route("/magazine")
def magazine():
    archive = []
    magazine_url = None

    for item in MAGAZINE_ISSUES:
        record = dict(item)
        raw_url = record.get("url")
        resolved_url = None
        if raw_url:
            if raw_url.startswith(("http://", "https://", "/")):
                resolved_url = raw_url
            else:
                resolved_url = url_for("static", filename=raw_url.lstrip("/"))
            record["url"] = resolved_url
            if magazine_url is None:
                magazine_url = resolved_url
        archive.append(record)

    return render_template("language.html", magazine_url=magazine_url, archive=archive)


@app.route("/language")
def legacy_language():
    return redirect(url_for("magazine"))


@app.route("/blogs")
def blogs():
    # Dedicated blogs listing page - attempt to extract full post content from magazine PDF
    posts = []
    # locate local magazine PDF if present
    mag_pdf = os.path.join(app.root_path, 'static', 'magazine', 'July-2025.pdf')
    for p in BLOG_POSTS:
        post = dict(p)
        # If URL contains a page anchor like '#page=9' try to extract that page from the PDF
        url = post.get('url', '') or ''
        m = re.search(r'page=(\d+)', url)
        if m and os.path.exists(mag_pdf):
            try:
                from pdfminer.high_level import extract_text
                page_num = int(m.group(1))
                # pdfminer uses 0-based page numbers
                text = extract_text(mag_pdf, page_numbers=[page_num - 1]) or ''
                # Basic cleanup
                text = text.strip()
                # Only overwrite `post['content']` from the PDF if the post has no
                # content defined in `BLOG_POSTS`. This ensures edits in `app.py`
                # are preserved and not clobbered by automatic PDF extraction.
                if text and not post.get('content'):
                    post['content'] = text
            except Exception:
                logging.exception('Failed to extract blog content from PDF')
        posts.append(post)

    return render_template("blogs.html", posts=posts, languages=LANGUAGES)


@app.route("/report_status")
def report_status():
    stats = db.get_stats()
    
    # Prepare JSON-safe data for JavaScript
    stats_json = {
        "reportsTimeSeries": stats.get("time_series", []),
        "ageData": [{"label": label, "value": count} for label, count in stats.get("age_ranges", {}).items()],
        "genderData": [
            {"label": "Female", "value": stats.get("gender", {}).get("female", 0)},
            {"label": "Male", "value": stats.get("gender", {}).get("male", 0)},
            {"label": "Other", "value": stats.get("gender", {}).get("other", 0)}
        ],
        "languagesData": stats.get("languages", []),
        "modalitiesData": stats.get("studies", []),
        "findingsData": [{"label": entry["label"].capitalize(), "value": entry["count"]} for entry in stats.get("diseases", [])]
    }
    
    return render_template("report_status.html", stats=stats, stats_json=stats_json)


@app.route("/payment")
def payment():
    # supply context expected by template
    structured_session = session.get("structured")
    if isinstance(structured_session, dict):
        structured = dict(structured_session)
    else:
        structured = {}

    structured.setdefault("report_type", "CT Scan")
    structured["price"] = f"{USD_PER_REPORT:.2f}"
    session["structured"] = structured

    kes_amount = USD_PER_REPORT * KES_PER_USD
    kes_display = f"{kes_amount:,.2f}".rstrip("0").rstrip(".")
    pricing = {
        "usd": USD_PER_REPORT,
        "usd_display": f"{USD_PER_REPORT:.2f}",
        "kes": kes_amount,
        "kes_display": kes_display,
        "tokens": TOKENS_PER_REPORT,
        "exchange_rate": KES_PER_USD,
    }
    lang = session.get("language", "English")
    return render_template("payment.html", structured=structured, language=lang, pricing=pricing)


@app.route("/help")
def help_page():
    return render_template("help.html")


@app.route("/team")
def team():
    """Team page with member bios and photos"""
    return render_template("team.html")


@app.route("/profile")
def profile():
    username = session.get("username")
    if not username:
        return redirect(url_for("login"))
    
    # Get user's feedback submissions
    user_feedback = db.get_user_feedback(username)
    
    is_admin = db.has_any_role(username, ["admin", "radiologist"])
    mfa_enabled = db.is_totp_enabled(username)
    return render_template("profile.html", feedback_list=user_feedback, is_admin=is_admin, mfa_enabled=mfa_enabled)


@app.route("/submit-feedback", methods=["POST"])
def submit_feedback():
    """Handle feedback submission from radiologists/users"""
    username = session.get("username")
    if not username:
        flash("Please log in to submit feedback.", "error")
        return redirect(url_for("login"))
    
    try:
        feedback_type = request.form.get("feedback_type", "").strip()
        subject = request.form.get("subject", "").strip()
        original_text = request.form.get("original_text", "").strip()
        corrected_text = request.form.get("corrected_text", "").strip()
        description = request.form.get("description", "").strip()
        
        if not feedback_type or not subject:
            flash("Please provide feedback type and subject.", "error")
            return redirect(url_for("profile"))
        
        feedback_id = db.submit_feedback(
            username=username,
            feedback_type=feedback_type,
            subject=subject,
            original=original_text,
            corrected=corrected_text,
            description=description
        )
        
        logging.info("Feedback #%d submitted (type=%s, subject_len=%d)", feedback_id, feedback_type, len(subject or ""))
        flash("Thank you! Your feedback has been submitted successfully.", "success")

    except Exception:
        logging.exception("Failed to submit feedback")
        flash("Sorry, there was an error submitting your feedback. Please try again.", "error")
    
    return redirect(url_for("profile"))


@app.route("/feedback-admin")
def feedback_admin():
    """Admin view to review all feedback submissions"""
    username = session.get("username")
    if not username:
        return redirect(url_for("login"))
    
    if not db.has_any_role(username, ["admin", "radiologist"]):
        flash("Access denied. Admin privileges required.", "error")
        return redirect(url_for("profile"))

    # Get filter status from query params
    status_filter = request.args.get("status", "pending")
    if status_filter == "all":
        all_feedback = db.get_all_feedback()
    else:
        all_feedback = db.get_all_feedback(status=status_filter)
    
    return render_template("feedback_admin.html", feedback_list=all_feedback, status_filter=status_filter)


@app.route("/review-feedback/<int:feedback_id>", methods=["POST"])
def review_feedback(feedback_id):
    """Admin action to approve/reject feedback"""
    username = session.get("username")
    if not username:
        return redirect(url_for("login"))
    
    if not db.has_any_role(username, ["admin", "radiologist"]):
        flash("Access denied. Admin privileges required.", "error")
        return redirect(url_for("feedback_admin"))
    
    try:
        status = request.form.get("status", "").strip()
        admin_notes = request.form.get("admin_notes", "").strip()
        
        if status not in ["approved", "rejected", "implemented"]:
            flash("Invalid status.", "error")
            return redirect(url_for("feedback_admin"))
        
        db.update_feedback_status(
            feedback_id=feedback_id,
            status=status,
            reviewed_by=username,
            admin_notes=admin_notes
        )
        
        logging.info("Feedback #%d reviewed by %s: %s", feedback_id, username, status)
        flash(f"Feedback #{feedback_id} marked as {status}.", "success")

    except Exception:
        logging.exception("Failed to review feedback")
        flash("Sorry, there was an error processing your request.", "error")
    
    return redirect(url_for("feedback_admin"))


@app.route("/contact-support", methods=["POST"])
def contact_support():
    """Handle contact support form submission"""
    try:
        subject = request.form.get("subject", "").strip()
        message = request.form.get("message", "").strip()

        # Log only lengths, never user-supplied free text (may contain PHI)
        logging.info("Support request received (subject_len=%d, message_len=%d)", len(subject), len(message))
        
        flash("Thank you for contacting us! We'll get back to you soon.", "success")
    except Exception:
        logging.exception("Failed to process support request")
        flash("Sorry, there was an error submitting your message. Please try again.", "error")
    
    return redirect(url_for("help_page"))


_LOGIN_RATE_LIMIT = os.getenv("LOGIN_RATE_LIMIT", "10 per 15 minutes")
_LOGIN_DUMMY_HASH = generate_password_hash("placeholder-do-not-use-as-credential")


def _login_rate_limit_decorator(view):
    return limiter.limit(_LOGIN_RATE_LIMIT, methods=["POST"])(view) if limiter else view


@app.route("/login", methods=["GET", "POST"])
@_login_rate_limit_decorator
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        # Lockout check (independent of whether the user exists, to avoid leaking existence)
        if username and db.is_account_locked(username):
            flash("This account is temporarily locked. Try again in a few minutes.", "error")
            logging.warning("Login attempt on locked account ip=%s", request.remote_addr)
            return render_template("login.html")

        user = db.get_user_by_username(username)

        # Always perform a password hash check to equalize timing between
        # "user exists" and "user does not exist" code paths.
        if user and user["password_hash"]:
            ok = check_password_hash(user["password_hash"], password)
        else:
            check_password_hash(_LOGIN_DUMMY_HASH, password)
            ok = False

        if ok:
            db.record_successful_login(username)
            session.clear()  # rotate session on auth change to mitigate fixation
            if db.is_totp_enabled(username):
                session["_mfa_pending"] = username
                session.permanent = True
                return redirect(url_for("mfa_verify"))
            session["username"] = username
            session["_last_active"] = time.time()
            session.permanent = True
            return redirect(url_for("dashboard"))

        # Generic message — never leak whether username or password was wrong
        if user:
            db.record_failed_login(username)
        flash("Invalid username or password.", "error")
        logging.warning("Failed login attempt ip=%s", request.remote_addr)
    return render_template("login.html")


@app.route("/login/google")
def login_google():
    redirect_uri = url_for('authorize', _external=True)
    return google.authorize_redirect(redirect_uri)


@app.route("/authorize")
def authorize():
    token = google.authorize_access_token()
    resp = google.get('userinfo')
    user_info = resp.json()
    email = user_info['email']
    google_id = user_info['id']
    name = user_info.get('name', email.split('@')[0])

    # Check if user exists
    user = db.get_user_by_google_id(google_id)
    if not user:
        # Create user if it doesn't exist
        # We use a unique username if the original one is taken
        username = name
        if db.get_user_by_username(username):
            username = f"{name}_{google_id[:5]}"
        
        db.create_oauth_user(username, email, google_id)
        user = db.get_user_by_google_id(google_id)

    if not user:
        abort(500)
    session.clear()  # rotate session to prevent fixation
    session["username"] = user["username"]
    flash("Logged in with Google successfully.", "success")
    return redirect(url_for("dashboard"))


_SIGNUP_RATE_LIMIT = os.getenv("SIGNUP_RATE_LIMIT", "10 per hour")


def _signup_rate_limit_decorator(view):
    return limiter.limit(_SIGNUP_RATE_LIMIT, methods=["POST"])(view) if limiter else view


_PASSWORD_MIN_LEN = int(os.getenv("PASSWORD_MIN_LENGTH", "12"))


def _password_complexity_error(password: str):
    """Return a human-readable error if the password is too weak, else None."""
    if not password or len(password) < _PASSWORD_MIN_LEN:
        return f"Password must be at least {_PASSWORD_MIN_LEN} characters long."
    classes = 0
    if any(c.islower() for c in password): classes += 1
    if any(c.isupper() for c in password): classes += 1
    if any(c.isdigit() for c in password): classes += 1
    if any(not c.isalnum() for c in password): classes += 1
    if classes < 3:
        return "Password must mix at least 3 of: lowercase, uppercase, digits, symbols."
    if password.lower() in {"password", "passw0rd", "qwerty12345!", "letmeinplease"}:
        return "Password is too common. Choose something less guessable."
    return None


@app.route("/signup", methods=["GET", "POST"])
@_signup_rate_limit_decorator
def signup():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        # Username constraints (length + charset)
        if not (3 <= len(username) <= 32) or not re.match(r"^[A-Za-z0-9_.\-]+$", username):
            flash("Username must be 3–32 characters using letters, digits, _ . or -.", "error")
            return render_template("signup.html")

        # Password policy
        pw_err = _password_complexity_error(password)
        if pw_err:
            flash(pw_err, "error")
            return render_template("signup.html")

        if db.get_user_by_username(username):
            flash("Username already exists. Please choose a different one.", "error")
        else:
            password_hash = generate_password_hash(password)
            db.create_user(username, password_hash)
            flash("Account created successfully. Please log in.", "success")
            return redirect(url_for("login"))
    return render_template("signup.html")


@app.route("/logout")
def logout():
    session.clear()  # Clear entire session instead of just username
    flash("Logged out successfully.", "success")
    return redirect(url_for("login"))


# --- MFA routes ----------------------------------------------------------

@app.route("/mfa/verify", methods=["GET", "POST"])
def mfa_verify():
    import pyotp
    pending_username = session.get("_mfa_pending")
    if not pending_username:
        return redirect(url_for("login"))

    if request.method == "POST":
        code = (request.form.get("code") or "").strip()
        secret = db.get_totp_secret(pending_username)
        if not secret:
            session.pop("_mfa_pending", None)
            flash("MFA not configured. Please contact support.", "error")
            return redirect(url_for("login"))

        if pyotp.TOTP(secret).verify(code, valid_window=1):
            session.pop("_mfa_pending", None)
            session["username"] = pending_username
            session["_last_active"] = time.time()
            session.permanent = True
            return redirect(url_for("dashboard"))

        flash("Invalid authentication code. Please try again.", "error")

    return render_template("mfa_verify.html")


@app.route("/mfa/setup", methods=["GET", "POST"])
def mfa_setup():
    import base64
    import pyotp
    import qrcode

    username = session.get("username")
    if not username:
        return redirect(url_for("login"))

    if request.method == "POST":
        code = (request.form.get("code") or "").strip()
        pending_secret = session.get("_mfa_setup_secret")
        if not pending_secret:
            flash("Setup session expired. Please try again.", "error")
            return redirect(url_for("mfa_setup"))

        if pyotp.TOTP(pending_secret).verify(code, valid_window=1):
            db.set_totp_secret(username, pending_secret)
            db.enable_totp(username)
            session.pop("_mfa_setup_secret", None)
            flash("Two-factor authentication enabled.", "success")
            return redirect(url_for("profile"))

        flash("Invalid code — please try again.", "error")

    secret = pyotp.random_base32()
    session["_mfa_setup_secret"] = secret
    uri = pyotp.TOTP(secret).provisioning_uri(name=username, issuer_name="Inside Imaging")
    buf = io.BytesIO()
    qrcode.make(uri).save(buf, kind="PNG")
    qr_b64 = base64.b64encode(buf.getvalue()).decode()

    return render_template("mfa_setup.html", secret=secret, qr_b64=qr_b64, uri=uri)


@app.route("/mfa/disable", methods=["POST"])
def mfa_disable():
    username = session.get("username")
    if not username:
        return redirect(url_for("login"))
    db.disable_totp(username)
    flash("Two-factor authentication disabled.", "success")
    return redirect(url_for("profile"))


# --- Data retention admin endpoint ---------------------------------------

@app.route("/admin/retention-purge", methods=["POST"])
@require_role("admin")
def retention_purge():
    from src.db.connection import execute, fetch_one
    username = session.get("username")
    dry_run = request.form.get("dry_run", "1") != "0"

    count_row = fetch_one(
        "SELECT count(*) FROM patients WHERE expires_at IS NOT NULL AND datetime(expires_at) < datetime('now')"
    )
    expired_count = count_row[0] if count_row else 0

    if dry_run:
        return jsonify({"dry_run": True, "would_purge": expired_count})

    execute(
        "DELETE FROM patients WHERE expires_at IS NOT NULL AND datetime(expires_at) < datetime('now')"
    )
    _audit("retention_purge", resource_type="patients", details=f"Purged {expired_count} expired patient records")
    logging.info("Retention purge by %s: deleted %d expired patient records", username, expired_count)
    return jsonify({"dry_run": False, "purged": expired_count})


@app.route("/dicom/upload", methods=["POST"])
def dicom_upload():
    if "username" not in session:
        return jsonify({"error": "Unauthorized"}), 401

    try:
        storage = get_storage()
    except Exception as exc:
        logging.exception("Storage backend init failed")
        return jsonify({"error": str(exc)}), 500

    username = session["username"]
    tenant_id = db.get_user_tenant(username)
    tenant_cfg = db.get_tenant_integration(tenant_id) or {}
    phi_mode = (tenant_cfg.get("phi_mode") or "passthrough").lower()
    files = request.files.getlist("files")
    uploaded = 0
    skipped = 0
    study_uids = []

    for upload in files:
        if not upload or not upload.filename:
            continue
        dcm_bytes = b""
        try:
            dcm_bytes = upload.read()
            meta = parse_dicom_metadata(dcm_bytes)
        except Exception:
            # Filenames can contain accession # or patient identifiers — log size only
            logging.warning("Skipping invalid DICOM (bytes=%d)", len(dcm_bytes))
            skipped += 1
            continue

        if not meta.get("study_uid") or not meta.get("series_uid") or not meta.get("sop_uid"):
            logging.warning("Skipping DICOM with missing UIDs (bytes=%d)", len(dcm_bytes or b""))
            skipped += 1
            continue

        # Sanitize the raw DICOM file per tenant PHI mode BEFORE storage.
        # If passthrough, this is a no-op and original bytes go to S3.
        try:
            sanitized_bytes = sanitize_dicom_bytes(dcm_bytes, phi_mode, tenant_salt=tenant_id)
        except Exception:
            logging.exception("DICOM sanitization failed for sop=%s; storing original", meta.get("sop_uid"))
            sanitized_bytes = dcm_bytes

        try:
            storage_key = store_dicom(
                storage, sanitized_bytes, tenant_id,
                meta["study_uid"], meta["series_uid"], meta["sop_uid"],
            )
        except Exception:
            logging.exception("Storage put failed for sop=%s", meta.get("sop_uid"))
            skipped += 1
            continue

        meta["username"] = username
        meta["tenant_id"] = tenant_id
        meta["s3_key"] = storage_key
        meta["storage_backend"] = storage.backend_name

        try:
            db.upsert_dicom_study(meta)
            db.upsert_dicom_series(meta)
            db.insert_dicom_instance(meta)
        except Exception:
            logging.exception("DB persistence failed for sop=%s", meta.get("sop_uid"))
            skipped += 1
            continue

        _audit(
            "dicom.store",
            resource_type="instance",
            resource_uid=meta["sop_uid"],
            details=f"study={meta['study_uid']} backend={storage.backend_name}",
        )
        uploaded += 1
        if meta["study_uid"] not in study_uids:
            study_uids.append(meta["study_uid"])

    return jsonify({"uploaded": uploaded, "skipped": skipped, "study_uids": study_uids})


@app.route("/dicom/studies", methods=["GET"])
def dicom_studies():
    if not session.get("username"):
        return redirect(url_for("login"))
    tenant_id = db.get_user_tenant(session["username"])
    studies = db.get_dicom_studies(tenant_id, username=session["username"])
    _audit("dicom.studies.list", resource_type="study", details=f"count={len(studies)}")
    return render_template("dicom_studies.html", studies=studies)


@app.route("/dicom/studies/<study_uid>", methods=["GET"])
def dicom_study_detail(study_uid):
    if not session.get("username"):
        return redirect(url_for("login"))
    tenant_id = db.get_user_tenant(session["username"])
    series_list = db.get_dicom_series(tenant_id, study_uid)
    for series in series_list:
        series["instances"] = db.get_dicom_instances(tenant_id, series["series_instance_uid"])
    _audit("dicom.study.view", resource_type="study", resource_uid=study_uid)
    return render_template("dicom_viewer.html", study_uid=study_uid, series_list=series_list)


@app.route("/dicom/frame/<sop_uid>", methods=["GET"])
def dicom_frame(sop_uid):
    if not session.get("username"):
        return jsonify({"error": "Unauthorized"}), 401

    tenant_id = db.get_user_tenant(session["username"])
    record = db.get_dicom_instance_by_uid(tenant_id, sop_uid)
    if not record:
        _audit("dicom.frame.view", resource_type="instance", resource_uid=sop_uid, outcome="not_found")
        abort(404)

    try:
        storage = get_storage()
        dcm_bytes = storage.get(record["s3_key"])
        png_bytes = extract_frame_as_png(dcm_bytes)
    except ValueError:
        _audit("dicom.frame.view", resource_type="instance", resource_uid=sop_uid, outcome="no_pixel_data")
        abort(404)
    except Exception:
        logging.exception("Frame extraction failed for %s", sop_uid)
        _audit("dicom.frame.view", resource_type="instance", resource_uid=sop_uid, outcome="error")
        abort(500)

    _audit("dicom.frame.view", resource_type="instance", resource_uid=sop_uid)
    resp = Response(png_bytes, mimetype="image/png")
    resp.headers["Cache-Control"] = "max-age=3600"
    return resp


@app.route("/dicom/instances/<sop_uid>/raw", methods=["GET"])
def dicom_instance_raw(sop_uid):
    if not session.get("username"):
        return jsonify({"error": "Unauthorized"}), 401

    tenant_id = db.get_user_tenant(session["username"])
    record = db.get_dicom_instance_by_uid(tenant_id, sop_uid)
    if not record:
        _audit("dicom.raw.download", resource_type="instance", resource_uid=sop_uid, outcome="not_found")
        abort(404)

    try:
        storage = get_storage()
        dcm_bytes = storage.get(record["s3_key"])
    except Exception:
        logging.exception("Raw DICOM fetch failed for %s", sop_uid)
        _audit("dicom.raw.download", resource_type="instance", resource_uid=sop_uid, outcome="error")
        abort(500)

    _audit("dicom.raw.download", resource_type="instance", resource_uid=sop_uid)
    resp = Response(dcm_bytes, mimetype="application/dicom")
    resp.headers["Content-Disposition"] = f'attachment; filename="{sop_uid}.dcm"'
    return resp


if __name__ == "__main__":
    # Local dev only. Production should run gunicorn (see Procfile).
    # Debug mode is hard-gated to non-production environments.
    app.run(debug=not _IS_PRODUCTION, host="127.0.0.1", port=int(os.getenv("PORT", "5000")))
